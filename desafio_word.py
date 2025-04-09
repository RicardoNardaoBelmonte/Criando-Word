from docx import Document
from docx.shared import Cm

documento = Document()
#adicionando o titulo
documento.add_heading('Relatório Mensal de Vendas', 0)
#adicionando paragrafo
documento.add_paragraph('Neste mes de agosto foram realizados um total de 10 vendas de notebooks.Segue em anexo a tabela com os dados de cada venda realizada')
#Adicionadno subtitulo
documento.add_heading('Vendas de Agosto', level=1)
#adicionando tabela
registros = [
    ['modelo 1', 'R$ 1500,00', '10/08/2023', 'Roberto', 'Americanas'],
    ['modelo 1', 'R$ 1500,00', '10/08/2023', 'Roberto', 'Americanas'],
    ['modelo 1', 'R$ 1500,00', '10/08/2023', 'Roberto', 'Americanas'],
    ['modelo 1', 'R$ 1500,00', '10/08/2023', 'Roberto', 'Americanas'],
    ['modelo 1', 'R$ 1500,00', '10/08/2023', 'Roberto', 'Americanas'],
    ['modelo 1', 'R$ 1500,00', '10/08/2023', 'Roberto', 'Americanas'],
    ['modelo 1', 'R$ 1500,00', '10/08/2023', 'Roberto', 'Americanas'],
    ['modelo 1', 'R$ 1500,00', '10/08/2023', 'Roberto', 'Americanas'],
    ['modelo 1', 'R$ 1500,00', '10/08/2023', 'Roberto', 'Americanas'],
    ['modelo 1', 'R$ 1500,00', '10/08/2023', 'Roberto', 'Americanas'],
]

tabela = documento.add_table(rows=1, cols=5)
cabeçalho = tabela.rows[0].cells
cabeçalho[0].text = 'Modelo'
cabeçalho[1].text = 'Preço'
cabeçalho[2].text = 'Data'
cabeçalho[3].text = 'Vendedor'
cabeçalho[4].text = 'Loja'

for modelo, preco, data, vendedor, loja in registros:
    linha_atual = tabela.add_row().cells
    linha_atual[0].text = modelo
    linha_atual[1].text = preco
    linha_atual[2].text = data
    linha_atual[3].text = vendedor
    linha_atual[4].text = loja

#adicionando paragrafo novamente
documento.add_paragraph('Para as vendas deste mês, o Roberto foi o funcionario com a maior quantidade de vendas diretas da tabela.')
#adicionando um paragrafo em negrito
documento.add_paragraph('O lucro gerado na loja americanas graças ao roberto foi de R$ 15000,00.').bold = True
documento.save('Relatorios_de_vendas.docx')